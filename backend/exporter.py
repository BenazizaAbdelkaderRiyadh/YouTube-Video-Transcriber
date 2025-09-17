import os
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

BASE_DIR = os.path.dirname(os.path.abspath(__file__)) 

def create_pdf(transcript, filename="output.pdf"):
    file_path = os.path.join(BASE_DIR, filename)

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(file_path)
    elements = []


    elements.append(Paragraph("📌 Transcriber Made by Benaziza Abdelkader Riyadh visit my GitHub https://github.com/BenazizaAbdelkaderRiyadh", styles['Heading2']))
    elements.append(Spacer(1, 12))

    elements.append(Paragraph(transcript.replace("\n", "<br/>"), styles['Normal']))

    doc.build(elements)
    return file_path

def create_docx(transcript, filename="output.docx"):
    file_path = os.path.join(BASE_DIR, filename)
    doc = Document()
    doc.add_heading("📌 Made by me", level=1)
    doc.add_paragraph(transcript)
    doc.save(file_path)

    return file_path
